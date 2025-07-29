import re
import tempfile
import os
from pathlib import Path
import uuid
import asyncio
from typing import Optional, Dict, Any
from urllib.parse import urlparse
import aiohttp
import aiofiles
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain.tools import BaseTool
from markdownify import markdownify as md
import mammoth


class DocxGeneratorTool(BaseTool):
    """Microsoft Word DOCX Generator Tool."""
    name: str = "generate_ms_word_document"
    description: str = "Use this tool for generating DOCX, provide text in markdown format with sections, headings."
    output_dir: str = None

    def __init__(self, output_dir=None):
        """Initialize the DOCX generator tool."""
        super().__init__()
        self.output_dir = output_dir

    def _run(self, markdown_text: str, filename: str = None) -> dict:
        """Generate a DOCX document from markdown text."""
        try:
            # Create a unique filename if not provided
            if not filename:
                filename = f"document_{uuid.uuid4().hex[:8]}.docx"
            elif not filename.lower().endswith('.docx'):
                filename = f"{filename}.docx"

            # Get output directory, use current directory if not specified
            output_dir = self.output_dir or '.'
            os.makedirs(output_dir, exist_ok=True)

            # Create full file path
            file_path = Path(output_dir) / filename

            # Process the markdown to handle any format issues
            processed_text = self._preprocess_markdown(markdown_text)

            # Convert markdown to DOCX
            self._markdown_to_docx(processed_text, file_path)

            return {
                "filename": filename,
                "file_path": str(file_path),
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "type": "docx"
            }
        except Exception as e:
            # If any error occurs, return information about the error
            return {
                "error": str(e),
                "original_markdown": markdown_text
            }

    def _preprocess_markdown(self, text):
        """Preprocess markdown to handle common issues."""
        # Replace placeholder variables with empty strings
        text = re.sub(r'\{[a-zA-Z0-9_]+\}', '', text)

        # Handle f-strings that weren't evaluated
        text = re.sub(r'f"""(.*?)"""', r'\1', text, flags=re.DOTALL)
        text = re.sub(r"f'''(.*?)'''", r'\1', text, flags=re.DOTALL)

        # Remove triple backticks and language indicators (common in code blocks)
        text = re.sub(r'```[a-zA-Z]*\n', '', text)
        text = re.sub(r'```', '', text)

        # Fix any heading issues (ensure space after #)
        text = re.sub(r'(#+)([^ \n])', r'\1 \2', text)

        return text

    def _markdown_to_docx(self, markdown_text, output_path):
        """Convert markdown text to a DOCX document."""
        # Create a new Document
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Split the markdown into lines for processing
        lines = markdown_text.split('\n')

        # Process each line
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Handle headings
            if line.startswith('#'):
                # Count the number of # to determine heading level
                level = 0
                while level < len(line) and line[level] == '#':
                    level += 1

                # Get the heading text
                heading_text = line[level:].strip()

                # Add the heading with appropriate style
                if level <= 9:  # Word supports heading levels 1-9
                    heading = doc.add_heading(heading_text, level=level)
                else:
                    # If level is beyond supported, default to level 9
                    heading = doc.add_heading(heading_text, level=9)

            # Handle bullet lists
            elif line.startswith('* ') or line.startswith('- '):
                text = line[2:].strip()
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(text)

            # Handle numbered lists
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line).strip()
                p = doc.add_paragraph()
                p.style = 'List Number'
                p.add_run(text)

            # Handle blockquotes
            elif line.startswith('> '):
                text = line[2:].strip()
                p = doc.add_paragraph()
                p.style = 'Quote'
                p.add_run(text)

            # Handle horizontal rules
            elif line == '---' or line == '***' or line == '___':
                doc.add_paragraph('_' * 50)

            # Handle paragraphs (including empty lines)
            else:
                # Skip completely empty lines
                if not line and i + 1 < len(lines) and not lines[i+1].strip():
                    i += 1
                    continue

                # Start a new paragraph
                p = doc.add_paragraph()

                # Add the text, handling bold and italic formatting
                text = line

                # Process basic markdown formatting
                # Bold: **text** or __text__
                bold_pattern = r'\*\*(.*?)\*\*|__(.*?)__'
                # Italic: *text* or _text_
                italic_pattern = r'\*(.*?)\*|_(.*?)_'

                # Find all formatting markers
                formatting_markers = []

                # Find bold text
                for match in re.finditer(bold_pattern, text):
                    start, end = match.span()
                    content = match.group(1) or match.group(2)
                    formatting_markers.append((start, end, content, 'bold'))

                # Find italic text
                for match in re.finditer(italic_pattern, text):
                    start, end = match.span()
                    content = match.group(1) or match.group(2)
                    formatting_markers.append((start, end, content, 'italic'))

                # Sort markers by start position
                formatting_markers.sort(key=lambda x: x[0])

                # Apply formatting
                if formatting_markers:
                    # Add text with formatting
                    current_pos = 0
                    for start, end, content, format_type in formatting_markers:
                        # Add text before the formatted part
                        if start > current_pos:
                            p.add_run(text[current_pos:start])

                        # Add the formatted text
                        run = p.add_run(content)
                        if format_type == 'bold':
                            run.bold = True
                        elif format_type == 'italic':
                            run.italic = True

                        current_pos = end

                    # Add any remaining text
                    if current_pos < len(text):
                        p.add_run(text[current_pos:])
                else:
                    # No formatting, add the entire line
                    p.add_run(text)

            i += 1

        # Add a table if markdown contains a table-like structure
        if '|' in markdown_text:
            self._try_add_tables(doc, markdown_text)

        # Save the document
        doc.save(output_path)

    def _try_add_tables(self, doc, markdown_text):
        """Try to extract and add tables from markdown text."""
        # Find potential table rows (lines containing |)
        table_lines = [line.strip() for line in markdown_text.split('\n')
                      if '|' in line and line.strip().startswith('|')]

        if len(table_lines) >= 2:  # Need at least header and separator
            # Add a section break before the table
            doc.add_paragraph()

            # Extract header row
            header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

            # Check for separator row (contains only -, |, :)
            if all(all(c in '-:|' for c in cell.strip()) for cell in table_lines[1].split('|')[1:-1]):
                data_rows = table_lines[2:]
            else:
                data_rows = table_lines[1:]

            # Create table
            table = doc.add_table(rows=len(data_rows)+1, cols=len(header_cells))
            table.style = 'Table Grid'

            # Add header
            for i, cell in enumerate(header_cells):
                table.cell(0, i).text = cell

            # Add data rows
            for row_idx, row in enumerate(data_rows):
                cells = [cell.strip() for cell in row.split('|')[1:-1]]
                for col_idx, cell in enumerate(cells):
                    if col_idx < len(header_cells):  # Ensure we don't exceed columns
                        table.cell(row_idx+1, col_idx).text = cell


class WordToMarkdownTool(BaseTool):
    """Converts a Word document to Markdown format by downloading it from a URL."""
    name: str = "word_to_markdown_tool"
    description: str = (
        "Converts a Word document to Markdown format from a URL. "
        "This tool downloads the Word document from the provided URL, "
        "converts it to Markdown format, and returns the content. "
        "Useful for processing Word documents and making them easier to analyze by LLMs."
        "\nThe input must be the complete URL of the Word document."
    )
    return_direct: bool = False
    _temp_dir: Optional[str] = None

    async def _download_file(self, url: str) -> str:
        """Downloads a file from a URL to a temporary file."""
        # Create a temporary directory if it doesn't exist
        if not self._temp_dir:
            self._temp_dir = tempfile.mkdtemp()

        # Get the filename from the URL
        parsed_url = urlparse(url)
        filename = os.path.basename(parsed_url.path)
        if not filename.endswith(('.docx', '.doc')):
            filename += '.docx'  # Add extension if it doesn't exist

        # Complete path to the temporary file
        file_path = os.path.join(self._temp_dir, filename)

        # Download the file
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise Exception(f"Error downloading the file: {response.status}")

                # Save the file
                async with aiofiles.open(file_path, 'wb') as f:
                    await f.write(await response.read())

        return file_path

    async def _convert_to_markdown(self, file_path: str) -> str:
        """Converts a Word document to Markdown."""
        # Use mammoth to convert to HTML and then to Markdown
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            markdown_text = md(html)

            # If there are warning messages, add them as a comment at the beginning
            if result.messages:
                warnings = "\n".join([f"<!-- Warning: {msg} -->" for msg in result.messages])
                markdown_text = f"{warnings}\n\n{markdown_text}"

            return markdown_text

    async def _process_word_document(self, url: str) -> Dict[str, Any]:
        """Processes a Word document from a URL and converts it to Markdown."""
        try:
            file_path = await self._download_file(url)
            markdown_text = await self._convert_to_markdown(file_path)

            # Cleanup of temporary files
            if os.path.exists(file_path):
                os.remove(file_path)

            return {
                "markdown": markdown_text,
                "source_url": url,
                "success": True
            }
        except Exception as e:
            return {
                "error": str(e),
                "source_url": url,
                "success": False
            }
        finally:
            # Ensure cleanup of the temporary directory if it's empty
            if self._temp_dir and os.path.exists(self._temp_dir) and not os.listdir(self._temp_dir):
                os.rmdir(self._temp_dir)

    async def _arun(self, url: str) -> Dict[str, Any]:
        """Runs the tool asynchronously."""
        return await self._process_word_document(url)

    def _run(self, url: str) -> Dict[str, Any]:
        """Runs the tool synchronously."""
        loop = asyncio.get_event_loop()
        if loop.is_running():
            return loop.run_until_complete(self._process_word_document(url))
        else:
            return asyncio.run(self._process_word_document(url))
