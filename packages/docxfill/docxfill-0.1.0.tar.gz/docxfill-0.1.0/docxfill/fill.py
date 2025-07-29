import os
import time
import zipfile
from io import BytesIO
from docx import Document
from docx.shared import Inches
import openpyxl
import ast

def convert_to_number_if_possible(value):
    """Convert literal to int or float if it's convertible."""
    try:
        evaluated = ast.literal_eval(value)
        if isinstance(evaluated, (int, float)):
            return evaluated
    except (ValueError, SyntaxError):
        pass
    return value

def replace_text_in_paragraph(paragraph, replace_text, replace_images):
    """Search and replace {{}} in a Word paragraph with text or images."""
    runs = paragraph.runs
    stack, start, segments = [], False, []

    for item in runs:
        text = item.text
        if "{{" in text:
            start = True
        if start:
            segments.append(item)
            stack.append(text)

        if "}}" in text and start:
            placeholder = "".join(stack).split("{{")[1].split("}}")[0]

            if placeholder in replace_text:
                segments[1].text = replace_text[placeholder]
                segments[0].text = segments[0].text.replace("{{", "")
                segments[-1].text = segments[-1].text.replace("}}", "")
                for s in segments[2:-1]:
                    s.text = ''

            if placeholder in replace_images:
                run = paragraph.add_run()
                run.add_picture(replace_images[placeholder], width=Inches(6))
                
                segments[-1].text = segments[-1].text.replace("}}", "")
                segments[0].text = segments[0].text.replace("{{", "")
                for s in segments[1:-1]:
                    s.text = ''

            stack, start, segments = [], False, []

def replace_text_in_excel_sheet(sheet, replace):
    """Search and replace {{}} in an Excel sheet's cells."""
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str) and '{{' in cell.value and '}}' in cell.value:
                placeholder = cell.value.split('{{')[1].split('}}')[0]
                if placeholder in replace:
                    cell.value = convert_to_number_if_possible(replace[placeholder])

def extract_and_modify_embedded_excel(docx_file, fname, replace):
    """Extract, modify, and return a modified .xls from a .docx."""
    with docx_file.open(fname) as emb_file:
        with BytesIO(emb_file.read()) as xls_file:
            wb = openpyxl.load_workbook(xls_file, data_only=True)
            for sheet in wb.sheetnames:
                replace_text_in_excel_sheet(wb[sheet], replace)

            modified = BytesIO()
            wb.save(modified)
            return modified.getvalue()

def check_unfill(paragraph):
    """Identify placeholders that were not filled in a Word paragraph."""
    runs = paragraph.runs
    stack, start, segments, unfills = [], False, [], []

    for item in runs:
        text = item.text
        if "{{" in text:
            start = True
        if start:
            segments.append(item)
            stack.append(text)

        if "}}" in text and start:
            placeholder = "".join(stack).split("{{")[1].split("}}")[0]
            unfills.append(placeholder)
            stack, start, segments = [], False, []

    return unfills

def extract_and_modify_docx(file_path, output_file, text, images):
    """Extract, modify, and pack back into a new .docx."""
    with zipfile.ZipFile(file_path, 'r') as src, \
         zipfile.ZipFile(output_file, 'w') as dst:

        for item in src.infolist():
            if item.filename.startswith('word/embeddings/') and item.filename.endswith('.xlsx'):
                modified = extract_and_modify_embedded_excel(src, item.filename, text)
                dst.writestr(item.filename, modified)
            else:
                dst.writestr(item.filename, src.read(item.filename))

    doc = Document(output_file)

    for para in doc.paragraphs:
        replace_text_in_paragraph(para, text, images)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, text, images)

    doc.save(output_file)

    doc = Document(output_file)
    unfills = []
    for para in doc.paragraphs:
        unfills += check_unfill(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    unfills += check_unfill(para)

    return unfills

def fill(file_path, output_file, text=None, images=None):
    """API main: execute filling directly from files."""
    text = text or {}
    images = images or {}
    unfills = extract_and_modify_docx(file_path, output_file, text, images)

    for k, v in images.items():
        text[k] = v

    return {
        "success": True,
        "output_file": output_file,
        "filled": {**text, **images},
        "unfilled": unfills
    }
