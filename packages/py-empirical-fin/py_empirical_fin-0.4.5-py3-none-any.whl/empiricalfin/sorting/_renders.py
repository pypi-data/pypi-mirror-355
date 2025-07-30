# docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# word template
import tempfile
import shutil
import importlib.resources as pkg_resources




class HtmlRenderer(object):
    def __init__(self, means, tstats, hhead, vhead, notes=[]):
        self.means, self.tstats = means, tstats
        self.hhead, self.vhead = hhead, vhead
        self.notes = notes

        self.hlabel, self.vlabel = means.columns, means.index

        self.table = '<table><tbody>'

    def cell(self, content="", bold=False):
        if bold:
            node = "<th>{}</th>"
        else:
            node = "<td>{}</td>"

        return node.format(content)

    def add_cells(self, series, bold=False, center=True):
        html = ""
        for value in series:
            html += self.cell(value, bold=bold)
        return html

    def _add_hhead(self):
        hhead = '<tr>' + self.cell() * 2  # skip two columns for vhead & vlabel
        hhead += f'<th colspan="{len(self.hlabel)}" style="text-align: center;">'  # hhead merge cells
        hhead += f'{self.hhead}</th></tr>'  # add hhead

        self.table += hhead

    def _add_hlabel(self):
        hlabel = '<tr>' + self.cell() * 2  # skip two columns for vhead & vlabel
        hlabel += self.add_cells(self.hlabel, bold=True) + '</tr>'  # add hlabel

        self.table += hlabel

    def _add_tbody(self):
        # iters rows and insert into html table
        for i in range(len(self.vlabel)):
            mrow = self.means.iloc[i]
            trow = self.tstats.iloc[i]

            if i == 0:  # first row add vhead
                # mean
                tbody = f'<tr><th rowspan="{len(self.vlabel) * 2}">'  # vhead merge cells
                tbody += f'{self.vhead}</th>'  # add vhead
                tbody += f'<th>{self.vlabel[i]}</th>'  # add vlabel
                tbody += self.add_cells(mrow) + '</tr>'  # add content
            else:
                # mean
                tbody += f"<tr><th>{self.vlabel[i]}</th>"  # add vlabel
                tbody += self.add_cells(mrow) + '</tr>'  # add content

            # tstats
            tbody += f"<tr><th></th>"  # skip vlabel cell
            tbody += self.add_cells(trow) + '</tr>'  # add content

        tbody += '</tbody>'  # close tbody tag
        self.table += tbody

    def add_notes(self):
        self.table += '<p style="font-size:12px; line-height:10px;">'
        self.table += 'Notes:'
        for i, note in enumerate(self.notes):
            self.table += '<p style="font-size:12px; line-height:4px;">'
            self.table += f'({i + 1}) {note}</p>'

    def render(self):
        self._add_hhead()
        self._add_hlabel()
        self._add_tbody()

        self.table += '</table>'  # close table tag
        self.add_notes()
        return self.table


class DocxRenderer(object):
    def __init__(self, means, tstats, hhead, vhead, notes=[], doc=None, template="template.docx",
                 style="academic table"):
        self.means, self.tstats = means, tstats
        self.hhead, self.vhead = hhead, vhead
        self.notes = notes
        self.style = style

        self.hlabel, self.vlabel = means.columns, means.index

        if doc:
            self.doc = doc
        else:
            with pkg_resources.open_binary('empiricalfin', 'template.docx') as file:
                file_content = file.read()

            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            self.doc = Document(temp_file_path)
            # 删除临时文件
            shutil.rmtree(temp_file_path, ignore_errors=True)

        self.table = self.doc.add_table(rows=2, cols=len(self.hlabel) + 2)

    @property
    def rows(self):
        return self.table.rows

    @property
    def cols(self):
        return self.table.columns

    def add_cells(self, series, cells):
        for text, cell in zip(series, cells):
            cell.text = text

    def _add_hhead(self):
        # merge cells for hhead
        row = self.rows[0]
        begin = row.cells[2]
        end = row.cells[-1]
        begin.merge(end)
        # add hhead
        begin.text = self.hhead

    def _add_hlabel(self):
        row = self.rows[1].cells[2:]
        for cell, text in zip(row, self.hlabel):
            cell.text = str(text)

    def _add_tbody(self):
        # each iteration add two rows, one for mean and another for tstats
        for i in range(len(self.vlabel)):
            mrow = self.means.iloc[i]
            trow = self.tstats.iloc[i]

            mcells = self.table.add_row().cells  # add row
            mcells[1].text = str(self.vlabel[i])  # add vlabel

            # add mean
            mcells = mcells[2:]
            self.add_cells(mrow, mcells)

            # add tstats
            tcells = self.table.add_row().cells[2:]
            self.add_cells(trow, tcells)

    def _add_vhead(self):
        col = self.cols[0]
        # merge cells for vhead
        begin = col.cells[2]
        end = col.cells[-1]
        begin.merge(end)
        # add vhead
        begin.text = self.vhead

    def _alignment(self):
        for row in self.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_note(self):
        self.doc.add_paragraph("Notes:")
        for i, note in enumerate(self.notes):
            note = f"({i + 1})" + note
            self.doc.add_paragraph(note)
        self.doc.add_paragraph("=" * 59)
        self.doc.add_paragraph(" ")

    def render(self):
        self._add_hhead()
        self._add_hlabel()
        self._add_tbody()
        self._add_vhead()
        self._alignment()
        self.table.style = self.style
        self._add_note()
        return self.doc